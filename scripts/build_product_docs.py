from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Interview-OS-product-spec-v0.5.0.md"
OUTPUT = ROOT / "docs" / "Interview-OS-产品与架构说明-v0.5.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1D2F44"
MUTED = "66778A"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EEF5FC"
GREEN_FILL = "E7F4EE"
GREEN = "176B53"
BORDER = "D8E0E8"


def set_run_font(run, size=None, color=None, bold=None, italic=None, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Product Title", "Product Subtitle", "Callout", "Code Block"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    title = styles["Product Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(5)
    subtitle = styles["Product Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(13.5)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)
    callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.line_spacing = 1.18
    code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string("334E68")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def configure_list_numbering(doc, style_name):
    """Apply compact_reference_guide list geometry to the built-in numbering."""
    style = doc.styles[style_name]
    num_pr = style._element.pPr.numPr if style._element.pPr is not None else None
    if num_pr is None or num_pr.numId is None:
        return
    num_id = str(num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    num_nodes = numbering.xpath(f'./w:num[@w:numId="{num_id}"]')
    if not num_nodes:
        return
    abstract_id = num_nodes[0].find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract_nodes = numbering.xpath(f'./w:abstractNum[@w:abstractNumId="{abstract_id}"]')
    if not abstract_nodes:
        return
    level_nodes = [
        child for child in abstract_nodes[0]
        if child.tag == qn("w:lvl") and child.get(qn("w:ilvl")) == "0"
    ]
    if not level_nodes:
        return
    level = level_nodes[0]
    p_pr = level.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        level.append(p_pr)
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    for child in list(tabs):
        tabs.remove(child)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "270")
    tabs.append(tab)
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")


def add_shading_to_paragraph(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first

    header_p = section.header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("INTERVIEW OS  |  产品与架构说明")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_p.add_run("\tv0.5.0")
    set_run_font(right, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer_p.add_run("Interview OS  ·  本地优先的一体化求职工作台  |  ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(footer_p)


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=9.5, color=DARK_BLUE, ascii_font="Consolas")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_bullet(doc, text, level=0, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)
    return p


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def table_widths(column_count):
    if column_count == 3:
        return [2100, 4800, 2460]
    if column_count == 4:
        return [1800, 2600, 1500, 3460]
    if column_count == 2:
        return [2500, 6860]
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, table_widths(len(rows[0])))
    set_table_borders(table)
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            if c_index >= 2 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value)
            for run in p.runs:
                set_run_font(run, size=9.2, bold=(r_index == 0), color=INK)
            if r_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
        if r_index == 0:
            set_repeat_table_header(table.rows[r_index])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_list_numbering(doc, "List Bullet")
    configure_list_numbering(doc, "List Number")
    configure_section(doc.sections[0], first=True)
    doc.core_properties.title = "Interview OS 产品与架构说明"
    doc.core_properties.subject = "产品需求文档与系统架构说明"
    doc.core_properties.author = "Interview OS Project"
    doc.core_properties.keywords = "Interview OS, PRD, 求职平台, 产品架构"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    kicker = p.add_run("PRODUCT SPECIFICATION  /  SYSTEM ARCHITECTURE")
    set_run_font(kicker, size=9.5, color=BLUE, bold=True)
    title = doc.add_paragraph("Interview OS", style="Product Title")
    subtitle = doc.add_paragraph("产品与架构说明", style="Product Subtitle")
    meta_rows = [
        ("版本", "v0.5.0 Career Workspace + Obsidian Phase 1"),
        ("阶段", "本地求职闭环与知识资产导出可用，真实外部连接器逐步接入前"),
        ("文档类型", "产品需求文档（PRD）+ 系统架构说明"),
        ("更新时间", "2026-07-21"),
        ("项目位置", r"D:\资料\release\IOS"),
    ]
    for label, value in meta_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=MUTED)
    doc.add_paragraph()
    status = doc.add_paragraph(style="Callout")
    add_inline(status, "当前结论：Interview OS 已形成岗位、简历、投递、训练、求职 Agent 和 Obsidian 知识资产导出的本地闭环。真实招聘平台抓取、外部推送、自动投递和双向同步仍保持关闭。")
    add_shading_to_paragraph(status, GREEN_FILL, GREEN)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Internal Product Document")
    set_run_font(r, size=9, color=MUTED, italic=True)

    doc.add_page_break()
    toc_title = doc.add_paragraph("文档目录", style="Heading 1")
    sections = [
        "1. 产品定义", "2. 目标用户与核心场景", "3. 端到端用户流程", "4. 产品信息架构",
        "5. 核心功能模块说明", "6. 数据与领域架构", "7. 连接器架构", "8. 系统技术架构",
        "9. 当前完成度", "10. 安全、隐私与风控", "11. 是否需要服务器与域名", "12. 已知问题",
        "13. 分阶段路线图", "14. 当前验收基线", "15. 相关文档", "16. 产品阶段结论",
    ]
    for item in sections:
        add_bullet(doc, item, numbered=False)
    doc.add_page_break()

    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    i = start
    in_code = False
    code_lines = []
    page_break_sections = {"## 5.", "## 7."}
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                p.add_run("\n".join(code_lines))
                add_shading_to_paragraph(p, LIGHT_FILL)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if stripped.startswith("## "):
            if any(stripped.startswith(prefix) for prefix in page_break_sections) and len(doc.paragraphs) > 5:
                doc.add_page_break()
            doc.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Callout")
            add_inline(p, stripped[2:])
            add_shading_to_paragraph(p, CALLOUT_FILL, BLUE)
        elif re.match(r"^\d+\.\s", stripped):
            add_bullet(doc, re.sub(r"^\d+\.\s", "", stripped), numbered=True)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped.replace("  ", " "))
        i += 1

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            if paragraph.style.name not in ("Code Block",):
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    qn("w:eastAsia"), "Microsoft YaHei"
                )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
